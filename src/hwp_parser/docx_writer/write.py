from __future__ import annotations

import io
import logging
from pathlib import Path
import zlib

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.shared import Inches

from hwp_parser.ir.convert import document_from_ir_json_file
from hwp_parser.ir.models import Document, DocumentSection, ImageBlock, Paragraph, Table, TableCell, TextRun

from .image_resolver import ImageResolutionContext, resolve_image_path
from .mapping import apply_character_style, map_alignment

LOGGER = logging.getLogger(__name__)


def write_docx(
    document: Document,
    output_path: Path,
    *,
    image_resolution_context: ImageResolutionContext | None = None,
) -> Path:
    docx_document = DocxDocument()

    _write_section(docx_document, document.header, section_name="header")
    _write_section(docx_document, document.footer, section_name="footer", is_footer=True)

    for block in document.blocks:
        if isinstance(block, Paragraph):
            _write_paragraph(docx_document, block)
            continue
        if isinstance(block, Table):
            _write_table(docx_document, block)
            continue
        if isinstance(block, ImageBlock):
            _write_image(docx_document, block, image_resolution_context=image_resolution_context)
            continue
        LOGGER.warning("Skipping unsupported block type in DOCX Phase 3: %r", type(block))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    docx_document.save(output_path)
    return output_path


def write_docx_from_ir_json(input_path: Path, output_path: Path) -> Path:
    document = document_from_ir_json_file(input_path)
    return write_docx(document, output_path)


def _write_section(
    docx_document: DocxDocument,
    section: DocumentSection | None,
    *,
    section_name: str,
    is_footer: bool = False,
) -> None:
    if section is None:
        return

    target = docx_document.sections[0].footer if is_footer else docx_document.sections[0].header
    if target.paragraphs and not target.paragraphs[0].text:
        paragraph = target.paragraphs[0]
        _clear_paragraph(paragraph)
        is_first = True
    else:
        is_first = False

    for block in section.blocks:
        if not isinstance(block, Paragraph):
            LOGGER.info(
                "Skipping non-paragraph %s block in DOCX Phase 1: %r",
                section_name,
                type(block),
            )
            continue
        paragraph = target.paragraphs[0] if is_first else target.add_paragraph()
        _fill_paragraph(paragraph, block)
        is_first = False


def _write_paragraph(docx_document: DocxDocument, paragraph_block: Paragraph) -> None:
    paragraph = docx_document.add_paragraph()
    _fill_paragraph(paragraph, paragraph_block)


def _write_table(docx_document: DocxDocument, table_block: Table) -> None:
    table = docx_document.add_table(rows=table_block.row_count, cols=table_block.column_count)
    covered: set[tuple[int, int]] = set()

    for row in table_block.rows:
        for cell_block in row.cells:
            coordinate = (cell_block.row_index, cell_block.column_index)
            if coordinate in covered:
                LOGGER.warning(
                    "Skipping overlapping table cell in DOCX Phase 2: row=%s col=%s",
                    cell_block.row_index,
                    cell_block.column_index,
                )
                continue

            docx_cell = table.cell(cell_block.row_index, cell_block.column_index)
            merged_cell = _merge_table_cell(table, docx_cell, cell_block)
            _fill_table_cell(merged_cell, cell_block)

            for covered_row in range(cell_block.row_index, cell_block.row_index + max(cell_block.rowspan, 1)):
                for covered_column in range(
                    cell_block.column_index,
                    cell_block.column_index + max(cell_block.colspan, 1),
                ):
                    covered.add((covered_row, covered_column))


def _write_image(
    docx_document: DocxDocument,
    image_block: ImageBlock,
    image_resolution_context: ImageResolutionContext | None = None,
) -> None:
    image_descriptor = _resolve_image_descriptor(
        image_block,
        image_resolution_context=image_resolution_context,
    )
    if image_descriptor is None:
        LOGGER.warning(
            "Skipping unresolved image block in DOCX Phase 3: binary_stream_ref=%s",
            image_block.binary_stream_ref,
        )
        return

    paragraph = docx_document.add_paragraph()
    run = paragraph.add_run()
    width_inches = _pixels_to_inches(image_block.width)
    height_inches = _pixels_to_inches(image_block.height)

    inline_shape = run.add_picture(
        image_descriptor,
        width=Inches(width_inches) if width_inches is not None else None,
        height=Inches(height_inches) if height_inches is not None else None,
    )
    _set_inline_shape_alt_text(inline_shape, image_block.alt_text)


def _fill_paragraph(paragraph: object, paragraph_block: Paragraph) -> None:
    alignment = map_alignment(paragraph_block.paragraph_style)
    if alignment is not None:
        paragraph.alignment = alignment

    appended = False
    for text_run in paragraph_block.text_runs:
        text = _phase1_text_for_run(text_run)
        if not text:
            continue
        run = paragraph.add_run(text)
        apply_character_style(run, text_run.character_style)
        appended = True

    if not appended and paragraph_block.text:
        run = paragraph.add_run(paragraph_block.text)
        if paragraph_block.text_runs:
            apply_character_style(run, paragraph_block.text_runs[0].character_style)


def _phase1_text_for_run(text_run: TextRun) -> str:
    if text_run.kind == "text":
        return text_run.text
    if text_run.kind == "field":
        if text_run.resolved_text:
            return text_run.resolved_text
        LOGGER.info("Skipping unresolved field run in DOCX Phase 1: field_type=%s", text_run.field_type)
        return ""
    LOGGER.info("Skipping unsupported run kind in DOCX Phase 1: kind=%s", text_run.kind)
    return ""


def _merge_table_cell(table: object, docx_cell: object, cell_block: TableCell) -> object:
    end_row = cell_block.row_index + max(cell_block.rowspan, 1) - 1
    end_column = cell_block.column_index + max(cell_block.colspan, 1) - 1

    if end_row == cell_block.row_index and end_column == cell_block.column_index:
        return docx_cell

    try:
        merged = docx_cell.merge(table.cell(end_row, end_column))
    except Exception as exc:
        LOGGER.warning(
            "Failed to merge DOCX table cell at row=%s col=%s rowspan=%s colspan=%s: %s",
            cell_block.row_index,
            cell_block.column_index,
            cell_block.rowspan,
            cell_block.colspan,
            exc,
        )
        if cell_block.colspan > 1:
            _apply_grid_span(docx_cell, cell_block.colspan)
        if cell_block.rowspan > 1:
            LOGGER.warning(
                "Vertical merge fallback not available in DOCX Phase 2: row=%s col=%s rowspan=%s",
                cell_block.row_index,
                cell_block.column_index,
                cell_block.rowspan,
            )
        return docx_cell

    if cell_block.colspan > 1:
        _apply_grid_span(merged, cell_block.colspan)
    return merged


def _fill_table_cell(docx_cell: object, cell_block: TableCell) -> None:
    if not docx_cell.paragraphs:
        paragraph = docx_cell.add_paragraph()
    else:
        paragraph = docx_cell.paragraphs[0]
    _clear_paragraph(paragraph)
    if cell_block.text:
        paragraph.add_run(cell_block.text)


def _apply_grid_span(docx_cell: object, colspan: int) -> None:
    tc_pr = docx_cell._tc.get_or_add_tcPr()
    grid_span = tc_pr.find(qn("w:gridSpan"))
    if grid_span is None:
        grid_span = tc_pr.makeelement(qn("w:gridSpan"))
        tc_pr.append(grid_span)
    grid_span.set(qn("w:val"), str(colspan))


def _clear_paragraph(paragraph: object) -> None:
    paragraph_element = paragraph._element
    for child in list(paragraph_element):
        paragraph_element.remove(child)


def _resolve_image_descriptor(
    image_block: ImageBlock,
    *,
    image_resolution_context: ImageResolutionContext | None = None,
) -> str | io.BytesIO | None:
    image_path = resolve_image_path(image_block, context=image_resolution_context)
    if image_path is not None:
        payload = image_path.read_bytes()
        if _is_supported_image_bytes(payload):
            return str(image_path)
        decompressed = _try_decompress_raw_deflate(payload)
        if decompressed is not None and _is_supported_image_bytes(decompressed):
            LOGGER.info(
                "Resolved compressed image payload for DOCX Phase 3: %s",
                image_block.binary_stream_ref,
            )
            return io.BytesIO(decompressed)
        LOGGER.warning(
            "Resolved image path is not a supported image payload: binary_stream_ref=%s path=%s",
            image_block.binary_stream_ref,
            image_path,
        )
    return None


def _pixels_to_inches(value: int | None) -> float | None:
    if value is None:
        return None
    return value / 96.0


def _set_inline_shape_alt_text(inline_shape: object, alt_text: str | None) -> None:
    if not alt_text:
        return
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", alt_text)


def _is_supported_image_bytes(payload: bytes) -> bool:
    return any(
        payload.startswith(signature)
        for signature in (
            b"\x89PNG\r\n\x1a\n",
            b"\xff\xd8\xff",
            b"GIF87a",
            b"GIF89a",
            b"BM",
        )
    )


def _try_decompress_raw_deflate(payload: bytes) -> bytes | None:
    try:
        return zlib.decompress(payload, -15)
    except zlib.error:
        return None
