from __future__ import annotations

import base64
from pathlib import Path
from zipfile import ZipFile

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

from hwp_parser.container import HwpContainerDumper
from hwp_parser.editor_model import ir_to_editor_model, write_docx_from_editor_model
from hwp_parser.ir.convert import document_from_debug_dir


PNG_BASE64 = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+a5u8AAAAASUVORK5CYII="
PNG_DATA_URL = f"data:image/png;base64,{PNG_BASE64}"


def _build_ir_document(sample_name: str, tmp_path: Path):
    summary = HwpContainerDumper(Path("hwp_samples") / sample_name).dump(tmp_path / "debug")
    return document_from_debug_dir(Path(summary.debug_dir))


def test_editor_model_export_writes_edited_paragraph_text_to_docx(tmp_path: Path) -> None:
    original = _build_ir_document("001_text_only.hwp", tmp_path)
    editor_model = ir_to_editor_model(original)
    editor_model["children"][0]["children"][0]["text"] = "수정된 제목"

    output_path = tmp_path / "edited.docx"
    write_docx_from_editor_model(editor_model, output_path, original_ir=original)

    docx_document = DocxDocument(output_path)
    assert docx_document.paragraphs[0].text == "수정된 제목"


def test_editor_model_export_writes_bold_font_size_and_alignment_to_docx(tmp_path: Path) -> None:
    original = _build_ir_document("002_paragraph_style.hwp", tmp_path)
    editor_model = ir_to_editor_model(original)
    paragraph = editor_model["children"][1]
    paragraph["attrs"]["alignment"] = "center"
    paragraph["children"][0]["marks"] = [
        {"type": "bold"},
        {"type": "fontSize", "value": 18},
    ]

    output_path = tmp_path / "styled.docx"
    write_docx_from_editor_model(editor_model, output_path, original_ir=original)

    docx_document = DocxDocument(output_path)
    exported_paragraph = docx_document.paragraphs[1]

    assert exported_paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert exported_paragraph.runs[0].bold is True
    assert exported_paragraph.runs[0].font.size is not None
    assert round(exported_paragraph.runs[0].font.size.pt, 1) == 18.0


def test_editor_model_export_writes_inserted_and_deleted_paragraphs_to_docx(tmp_path: Path) -> None:
    original = _build_ir_document("001_text_only.hwp", tmp_path)
    editor_model = ir_to_editor_model(original)

    editor_model["children"].insert(
        1,
        {
            "type": "paragraph",
            "id": "p999",
            "attrs": {"alignment": "left"},
            "children": [{"type": "text", "id": "text999", "text": "새 문단", "marks": []}],
        },
    )
    editor_model["children"] = [child for child in editor_model["children"] if child.get("id") != "p3"]

    output_path = tmp_path / "insert-delete.docx"
    write_docx_from_editor_model(editor_model, output_path, original_ir=original)

    docx_document = DocxDocument(output_path)
    paragraph_texts = [paragraph.text for paragraph in docx_document.paragraphs]

    assert "새 문단" in paragraph_texts
    assert "첫 번째 문단입니다" not in paragraph_texts


def test_editor_model_export_preserves_unchanged_table_and_image_blocks_safely(tmp_path: Path) -> None:
    original = _build_ir_document("008_mixed.hwp", tmp_path)
    editor_model = ir_to_editor_model(original)
    editor_model["children"][0]["children"][0]["text"] = "수정된 문서 제목"

    output_path = tmp_path / "mixed.docx"
    write_docx_from_editor_model(editor_model, output_path, original_ir=original)

    docx_document = DocxDocument(output_path)

    assert docx_document.paragraphs[0].text == "수정된 문서 제목"
    assert len(docx_document.tables) == 1
    assert [[cell.text for cell in row.cells] for row in docx_document.tables[0].rows] == [
        ["A", "B", "C"],
        ["D", "E", "F"],
    ]


def test_editor_model_export_writes_edited_table_cell_text_to_docx(tmp_path: Path) -> None:
    original = _build_ir_document("003_table_basic.hwp", tmp_path)
    editor_model = ir_to_editor_model(original)
    table_node = next(child for child in editor_model["children"] if child["type"] == "table")
    table_node["rows"][0]["cells"][1]["children"][0]["children"][0]["text"] = "수정된 셀"

    output_path = tmp_path / "table-edit.docx"
    write_docx_from_editor_model(editor_model, output_path, original_ir=original)

    docx_document = DocxDocument(output_path)
    assert [[cell.text for cell in row.cells] for row in docx_document.tables[0].rows] == [
        ["A", "수정된 셀", "C"],
        ["D", "E", "F"],
    ]


def test_editor_model_export_writes_replaced_image_to_docx(tmp_path: Path) -> None:
    original = _build_ir_document("004_image_basic.hwp", tmp_path)
    editor_model = ir_to_editor_model(original)
    image_node = next(child for child in editor_model["children"] if child["type"] == "image")
    image_node["attrs"]["src"] = PNG_DATA_URL
    image_node["attrs"]["alt"] = "교체된 이미지"

    output_path = tmp_path / "image-replaced.docx"
    write_docx_from_editor_model(editor_model, output_path, original_ir=original)

    with ZipFile(output_path) as archive:
        media_names = [name for name in archive.namelist() if name.startswith("word/media/")]
        assert media_names
        media_payloads = [archive.read(name) for name in media_names]

    assert base64.b64decode(PNG_BASE64) in media_payloads


def test_editor_model_export_writes_numbered_and_bullet_paragraph_styles_to_docx(tmp_path: Path) -> None:
    original = _build_ir_document("002_paragraph_style.hwp", tmp_path)
    editor_model = ir_to_editor_model(original)

    editor_model["children"][0]["attrs"]["listKind"] = "numbered"
    editor_model["children"][0]["attrs"]["listLevel"] = 0
    editor_model["children"][1]["attrs"]["listKind"] = "bullet"
    editor_model["children"][1]["attrs"]["listLevel"] = 0
    editor_model["children"][2]["attrs"]["listKind"] = "none"

    output_path = tmp_path / "lists.docx"
    write_docx_from_editor_model(editor_model, output_path, original_ir=original)

    docx_document = DocxDocument(output_path)

    assert docx_document.paragraphs[0].style.name == "List Number"
    assert docx_document.paragraphs[1].style.name == "List Bullet"
    assert docx_document.paragraphs[2].style.name != "List Number"
