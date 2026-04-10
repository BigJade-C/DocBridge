from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

from hwp_parser.container import HwpContainerDumper
from hwp_parser.ir.convert import document_from_debug_dir
from hwp_parser.docx_writer.write import write_docx


def _build_ir_document(sample_name: str, tmp_path: Path):
    summary = HwpContainerDumper(Path("hwp_samples") / sample_name).dump(tmp_path / "debug")
    return document_from_debug_dir(Path(summary.debug_dir))


def test_docx_writer_preserves_text_only_paragraph_text_order(tmp_path: Path) -> None:
    document = _build_ir_document("001_text_only.hwp", tmp_path)
    output_path = tmp_path / "001_text_only.docx"

    write_docx(document, output_path)

    docx_document = DocxDocument(output_path)
    paragraph_texts = [paragraph.text for paragraph in docx_document.paragraphs]

    assert paragraph_texts == ["제목입니다", "", "첫 번째 문단입니다", "", "두 번째 문단입니다"]


def test_docx_writer_reflects_paragraph_style_sample_formatting(tmp_path: Path) -> None:
    document = _build_ir_document("002_paragraph_style.hwp", tmp_path)
    output_path = tmp_path / "002_paragraph_style.docx"

    write_docx(document, output_path)

    docx_document = DocxDocument(output_path)
    title = docx_document.paragraphs[0]

    assert [paragraph.text for paragraph in docx_document.paragraphs] == [
        "제목입니다",
        "왼쪽 정렬 문단입니다",
        "오른쪽 정렬 문단입니다",
    ]
    assert title.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert title.runs[0].bold is True
    assert title.runs[0].font.size is not None
    assert round(title.runs[0].font.size.pt, 1) == 13.0


def test_docx_writer_writes_only_paragraph_blocks_for_mixed_sample(tmp_path: Path) -> None:
    document = _build_ir_document("008_mixed.hwp", tmp_path)
    output_path = tmp_path / "008_mixed.docx"

    write_docx(document, output_path)

    docx_document = DocxDocument(output_path)
    paragraph_texts = [paragraph.text for paragraph in docx_document.paragraphs]

    assert paragraph_texts == [
        "문서 제목",
        "첫 번째 문단입니다",
        "두 번째 문단입니다",
        "",
        "세 번째 문단입니다",
    ]
    assert len(docx_document.inline_shapes) == 1
