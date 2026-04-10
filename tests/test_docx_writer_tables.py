from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument

from hwp_parser.container import HwpContainerDumper
from hwp_parser.ir.convert import document_from_debug_dir
from hwp_parser.docx_writer.write import write_docx


def _build_ir_document(sample_name: str, tmp_path: Path):
    summary = HwpContainerDumper(Path("hwp_samples") / sample_name).dump(tmp_path / "debug")
    return document_from_debug_dir(Path(summary.debug_dir))


def _body_sequence(docx_document: DocxDocument) -> list[str]:
    sequence: list[str] = []
    for child in docx_document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            sequence.append("paragraph")
        elif child.tag.endswith("}tbl"):
            sequence.append("table")
    return sequence


def test_docx_writer_writes_basic_table_content(tmp_path: Path) -> None:
    document = _build_ir_document("003_table_basic.hwp", tmp_path)
    output_path = tmp_path / "003_table_basic.docx"

    write_docx(document, output_path)

    docx_document = DocxDocument(output_path)
    assert _body_sequence(docx_document) == ["paragraph", "table", "paragraph"]
    assert len(docx_document.tables) == 1

    table = docx_document.tables[0]
    assert len(table.rows) == 2
    assert len(table.columns) == 3
    assert [[cell.text for cell in row.cells] for row in table.rows] == [
        ["A", "B", "C"],
        ["D", "E", "F"],
    ]


def test_docx_writer_preserves_merged_table_structure(tmp_path: Path) -> None:
    document = _build_ir_document("007_table_merge.hwp", tmp_path)
    output_path = tmp_path / "007_table_merge.docx"

    write_docx(document, output_path)

    docx_document = DocxDocument(output_path)
    assert _body_sequence(docx_document) == ["paragraph", "table", "paragraph"]
    table = docx_document.tables[0]

    assert table.cell(0, 0).text == "A"
    assert table.cell(0, 1).text == "B+C"
    assert table.cell(1, 0).text == "D"
    assert table.cell(1, 1).text == "E"
    assert table.cell(1, 2).text == "F"

    top_middle_tc = table.cell(0, 1)._tc
    grid_span = top_middle_tc.tcPr.gridSpan
    assert grid_span is not None
    assert grid_span.val == 2


def test_docx_writer_preserves_table_position_in_mixed_sample(tmp_path: Path) -> None:
    document = _build_ir_document("008_mixed.hwp", tmp_path)
    output_path = tmp_path / "008_mixed_with_table.docx"

    write_docx(document, output_path)

    docx_document = DocxDocument(output_path)

    assert _body_sequence(docx_document) == [
        "paragraph",
        "paragraph",
        "table",
        "paragraph",
        "paragraph",
        "paragraph",
    ]
    assert len(docx_document.tables) == 1
    assert len(docx_document.inline_shapes) == 1
    assert [paragraph.text for paragraph in docx_document.paragraphs] == [
        "문서 제목",
        "첫 번째 문단입니다",
        "두 번째 문단입니다",
        "",
        "세 번째 문단입니다",
    ]
    assert [[cell.text for cell in row.cells] for row in docx_document.tables[0].rows] == [
        ["A", "B", "C"],
        ["D", "E", "F"],
    ]
