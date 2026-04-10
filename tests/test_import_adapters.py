from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

from hwp_parser.container import HwpContainerDumper
from hwp_parser.docx_writer.write import write_docx
from hwp_parser.importers.dispatch import import_document_to_ir_dict
from hwp_parser.ir.convert import document_from_debug_dir


def _build_ir_document(sample_name: str, tmp_path: Path):
    summary = HwpContainerDumper(Path("hwp_samples") / sample_name).dump(tmp_path / "debug")
    return document_from_debug_dir(Path(summary.debug_dir))


def _roundtrip_docx_payload(sample_name: str, tmp_path: Path) -> dict[str, object]:
    document = _build_ir_document(sample_name, tmp_path)
    docx_path = tmp_path / f"{Path(sample_name).stem}.docx"
    write_docx(document, docx_path)
    return import_document_to_ir_dict(
        docx_path,
        artifact_root=tmp_path / "imports",
    )


def test_hwp_import_adapter_uses_parser_pipeline(tmp_path: Path) -> None:
    payload = import_document_to_ir_dict(
        Path("hwp_samples/001_text_only.hwp"),
        artifact_root=tmp_path / "imports",
    )

    paragraph_texts = [
        block["text"]
        for block in payload["blocks"]
        if block["block_type"] == "paragraph" and block["text"]
    ]

    assert paragraph_texts == ["제목입니다", "첫 번째 문단입니다", "두 번째 문단입니다"]


def test_docx_import_adapter_reads_generated_docx_into_ir(tmp_path: Path) -> None:
    payload = _roundtrip_docx_payload("008_mixed.hwp", tmp_path)

    block_types = [block["block_type"] for block in payload["blocks"]]
    assert "paragraph" in block_types
    assert "table" in block_types
    assert "image" in block_types

    image = next(block for block in payload["blocks"] if block["block_type"] == "image")
    assert image["raw"]["binary_output_path"].endswith(".png")
    assert image["width"] == 444
    assert image["height"] == 517

    table = next(block for block in payload["blocks"] if block["block_type"] == "table")
    assert table["row_count"] == 2
    assert table["column_count"] == 3
    assert [cell["text"] for cell in table["rows"][0]["cells"]] == ["A", "B", "C"]
    assert [cell["text"] for cell in table["rows"][1]["cells"]] == ["D", "E", "F"]


def test_docx_import_adapter_preserves_paragraph_alignment_and_character_style(tmp_path: Path) -> None:
    payload = _roundtrip_docx_payload("002_paragraph_style.hwp", tmp_path)

    paragraphs = [block for block in payload["blocks"] if block["block_type"] == "paragraph" and block["text"]]
    assert [paragraph["text"] for paragraph in paragraphs] == [
        "제목입니다",
        "왼쪽 정렬 문단입니다",
        "오른쪽 정렬 문단입니다",
    ]
    assert paragraphs[0]["paragraph_style"]["alignment"] == "center"
    assert paragraphs[1]["paragraph_style"]["alignment"] == "left"
    assert paragraphs[2]["paragraph_style"]["alignment"] == "right"
    assert paragraphs[0]["text_runs"][0]["character_style"]["bold"] is True
    assert round(paragraphs[0]["text_runs"][0]["character_style"]["font_size_pt"], 1) == 13.0


def test_docx_import_adapter_preserves_header_and_footer_content(tmp_path: Path) -> None:
    payload = _roundtrip_docx_payload("006_header_footer.hwp", tmp_path)

    assert payload["header"]["blocks"][0]["text"] == "문서 헤더"
    assert payload["footer"]["blocks"][0]["text"] == "페이지 1"
    assert payload["blocks"][0]["text"] == "본문 내용입니다"


def test_docx_import_adapter_preserves_merged_table_structure(tmp_path: Path) -> None:
    payload = _roundtrip_docx_payload("007_table_merge.hwp", tmp_path)

    table = next(block for block in payload["blocks"] if block["block_type"] == "table")
    assert table["row_count"] == 2
    assert table["column_count"] == 3
    assert len(table["rows"][0]["cells"]) == 2
    assert table["rows"][0]["cells"][0]["text"] == "A"
    assert table["rows"][0]["cells"][1]["text"] == "B+C"
    assert table["rows"][0]["cells"][1]["colspan"] == 2
    assert [cell["text"] for cell in table["rows"][1]["cells"]] == ["D", "E", "F"]


def test_docx_import_adapter_detects_numbering_metadata(tmp_path: Path) -> None:
    docx_path = tmp_path / "numbering.docx"
    docx_document = DocxDocument()
    first = docx_document.add_paragraph("첫 번째 항목", style="List Number")
    first.alignment = WD_ALIGN_PARAGRAPH.LEFT
    docx_document.add_paragraph("두 번째 항목", style="List Number")
    docx_document.save(docx_path)

    payload = import_document_to_ir_dict(
        docx_path,
        artifact_root=tmp_path / "imports",
    )

    paragraphs = [block for block in payload["blocks"] if block["block_type"] == "paragraph"]
    assert [paragraph["text"] for paragraph in paragraphs] == ["첫 번째 항목", "두 번째 항목"]
    for paragraph in paragraphs:
        assert paragraph["list_info"] is not None
        assert paragraph["list_info"]["kind"] == "numbered"
        assert paragraph["list_info"]["raw"]["style_name"] == "List Number"
