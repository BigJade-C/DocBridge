from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument

from hwp_parser.container import HwpContainerDumper
from hwp_parser.ir.convert import document_from_debug_dir
from hwp_parser.docx_writer.write import write_docx


def _build_ir_document(sample_name: str, tmp_path: Path):
    summary = HwpContainerDumper(Path("hwp_samples") / sample_name).dump(tmp_path / "debug")
    return document_from_debug_dir(Path(summary.debug_dir))


def _body_sequence(docx_document: DocxDocument) -> list[str]:
    sequence: list[str] = []
    for child in docx_document.element.body.iterchildren():
        if child.tag.endswith("}tbl"):
            sequence.append("table")
            continue
        if child.tag.endswith("}p"):
            has_drawing = any(
                element.tag.endswith("}drawing")
                for element in child.iter()
            )
            sequence.append("image" if has_drawing else "paragraph")
    return sequence


def test_docx_writer_exports_image_when_binary_is_resolvable(tmp_path: Path) -> None:
    document = _build_ir_document("004_image_basic.hwp", tmp_path)
    output_path = tmp_path / "004_image_basic.docx"

    write_docx(document, output_path)

    docx_document = DocxDocument(output_path)
    assert len(docx_document.inline_shapes) == 1
    image_paragraph = docx_document.paragraphs[1]
    assert image_paragraph.text == ""
    assert _body_sequence(docx_document) == ["paragraph", "image"]


def test_docx_writer_preserves_mixed_document_order_with_image(tmp_path: Path) -> None:
    document = _build_ir_document("008_mixed.hwp", tmp_path)
    output_path = tmp_path / "008_mixed_with_image.docx"

    write_docx(document, output_path)

    docx_document = DocxDocument(output_path)

    assert _body_sequence(docx_document) == [
        "paragraph",
        "paragraph",
        "table",
        "paragraph",
        "image",
        "paragraph",
    ]
    assert len(docx_document.inline_shapes) == 1
    assert [paragraph.text for paragraph in docx_document.paragraphs] == [
        "문서 제목",
        "첫 번째 문단입니다",
        "두 번째 문단입니다",
        "",
        "세 번째 문단입니다",
    ]
